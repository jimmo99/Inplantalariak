#!/usr/bin/env python3
# -*- coding: utf-8 -*-

r"""
DOCX Extractor (GUI) — V8
- Mantiene estilos originales (no crea títulos artificiales).
- Agrupa por jerarquía: pega el heading padre REAL una sola vez.
- Portada original (sin numeración). Numeración empieza en 1 tras la portada.
- Opción GUI: Insertar Índice (TOC) tras la portada (excluye portada y el propio TOC).
- Barras de progreso + sanitizado de selección + 'sectPr' garantizado en cada recorte.

Requisitos:
    pip install python-docx docxcompose
"""

import os
import shutil
import tempfile
import traceback
from typing import List, Dict, Optional
from copy import deepcopy

import tkinter as tk
from tkinter import ttk, filedialog, messagebox

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_BREAK
from docxcompose.composer import Composer


# -------------------------- CONFIGURACIÓN RUTA -------------------------- #

LAST_PATH_FILE = "last_path.txt"

def leer_ultima_ruta():
    if os.path.isfile(LAST_PATH_FILE):
        with open(LAST_PATH_FILE, "r", encoding="utf-8") as f:
            ruta = f.read().strip()
            if os.path.isdir(ruta):
                return ruta
    return None

def guardar_ruta(ruta):
    with open(LAST_PATH_FILE, "w", encoding="utf-8") as f:
        f.write(ruta)



# -------------------------- UTILIDADES -------------------------- #

def get_heading_level(style_name: Optional[str]) -> Optional[int]:
    """Devuelve el nivel de 'Heading X' / 'Título X' si corresponde; si no, None."""
    if not style_name:
        return None
    name = style_name.strip().lower()
    candidates = ["heading ", "heading", "título ", "titulo ", "encabezado "]
    for p in candidates:
        if name.startswith(p):
            tail = name[len(p):].strip()
            digits = ""
            for ch in tail:
                if ch.isdigit():
                    digits += ch
                else:
                    break
            if digits:
                try:
                    return int(digits)
                except ValueError:
                    return None
    return None


# ---------------------- DETECCIÓN DE SECCIONES ---------------------- #

def detect_sections(input_path: str, levels: List[int]) -> List[Dict]:
    r"""
    Devuelve lista de SECCIONES (solo para los niveles seleccionados), cada una con:
      - title: str
      - level: int
      - parents: List[{'level': int, 'title': str, 'index': int}]  # cadena de ancestros reales
      - start_index: int
      - end_index: int
    """
    doc = Document(input_path)
    n_paragraphs = len(doc.paragraphs)

    # Pila: último heading visto por nivel → (title, index)
    stack = {}  # level -> (title, index)

    starts = []
    for i, p in enumerate(doc.paragraphs):
        style = getattr(p, "style", None)
        lvl_any = None
        if style is not None:
            lvl_any = get_heading_level(getattr(style, "name", None)) or \
                      get_heading_level(getattr(style, "style_id", "") or "")
        if lvl_any is not None:
            # limpiar niveles más profundos
            for k in list(stack.keys()):
                if k >= lvl_any:
                    del stack[k]
            title = (p.text or "").strip()
            stack[lvl_any] = (title, i)

            if lvl_any in levels:
                parents = []
                for lvl in sorted(stack.keys()):
                    if lvl < lvl_any:
                        t, idxp = stack[lvl]
                        parents.append({"level": lvl, "title": t, "index": idxp})
                starts.append({
                    "title": title,
                    "level": lvl_any,
                    "parents": parents,
                    "start_index": i
                })

    sections = []
    if not starts:
        sections.append({
            "title": "Documento completo",
            "level": 1,
            "parents": [],
            "start_index": 0,
            "end_index": n_paragraphs
        })
        return sections

    for idx, s in enumerate(starts):
        start_i = s["start_index"]
        end_i = starts[idx + 1]["start_index"] if idx + 1 < len(starts) else n_paragraphs
        sections.append({
            "title": s["title"],
            "level": s["level"],
            "parents": s["parents"],
            "start_index": start_i,
            "end_index": end_i
        })
    return sections


# ---------------------- RECORTE Y ESTRUCTURA DOCX ---------------------- #

def remove_outside_range(doc: Document, start_idx: int, end_idx_excl: int) -> None:
    """Elimina del body todo lo que esté fuera de [start_idx, end_idx_excl) por índice de párrafo."""
    if start_idx >= end_idx_excl:
        return
    p_count = 0
    body = doc.element.body
    to_remove = []
    for elm in list(body.iterchildren()):
        if elm.tag == qn('w:p'):
            in_range = (start_idx <= p_count < end_idx_excl)
            if not in_range:
                to_remove.append(elm)
            p_count += 1
        else:
            in_range = (start_idx <= p_count < end_idx_excl)
            if not in_range:
                to_remove.append(elm)
    for elm in to_remove:
        body.remove(elm)


def _ensure_trailing_sectpr(doc: Document) -> None:
    """Asegura que haya EXACTAMENTE un w:sectPr al final del body."""
    body = doc.element.body
    children = list(body.iterchildren())
    sect_elems = [elm for elm in children if elm.tag == qn('w:sectPr')]

    if not sect_elems:
        default = Document()
        default_sect = deepcopy(default.sections[0]._sectPr)  # CT_SectPr en la mayoría de versiones
        if not list(body.iterchildren()):
            body.append(OxmlElement('w:p'))
        body.append(default_sect)
        return

    last_sect = sect_elems[-1]
    for s in sect_elems[:-1]:
        body.remove(s)
    if list(body.iterchildren())[-1] is not last_sect:
        body.remove(last_sect)
        body.append(last_sect)


def _ensure_has_content(doc: Document) -> None:
    """Garantiza que exista al menos un párrafo o tabla ANTES del sectPr final."""
    body = doc.element.body
    children = list(body.iterchildren())
    has_content = any(elm.tag in (qn('w:p'), qn('w:tbl')) for elm in children if elm.tag != qn('w:sectPr'))
    if not has_content:
        body.insert(0, OxmlElement('w:p'))


def _ensure_valid_doc_structure(doc: Document) -> None:
    _ensure_trailing_sectpr(doc)
    _ensure_has_content(doc)


def slice_docx_by_paragraph_range(source_path: str, start_idx: int, end_idx_excl: int) -> str:
    """Crea un DOCX temporal con SOLO el rango indicado y estructura válida."""
    if start_idx < 0 or end_idx_excl <= start_idx:
        raise ValueError(f"Rango inválido: start={start_idx}, end={end_idx_excl}")
    tmp_dir = tempfile.mkdtemp(prefix="slice_")
    tmp_path = os.path.join(tmp_dir, os.path.basename(source_path))
    shutil.copy2(source_path, tmp_path)
    d = Document(tmp_path)
    remove_outside_range(d, start_idx, end_idx_excl)
    _ensure_valid_doc_structure(d)
    d.save(tmp_path)
    return tmp_path  # limpiar directorio lo hace el que llama


# ---------------------- TOC OPCIONAL + NUMERACIÓN ---------------------- #

def create_bookmark_doc(name: str, start: bool = True) -> Document:
    r"""Inserta un párrafo con bookmarkStart/End (para delimitar el TOC: \b <start> / \e <end>)."""
    d = Document()
    p = d.add_paragraph()
    r = p.add_run()
    if start:
        bk = OxmlElement('w:bookmarkStart'); bk.set(qn('w:id'), '0'); bk.set(qn('w:name'), name)
    else:
        bk = OxmlElement('w:bookmarkEnd');   bk.set(qn('w:id'), '0')
    r._r.append(bk)
    p.add_run("")
    return d


def create_toc_doc_ranged(start_bmk: str, end_bmk: str, max_level: int = 9, title: str = "Índice") -> Document:
    r"""Documento con heading y campo TOC limitado al rango de marcadores."""
    d = Document()
    d.add_heading(title, level=1)
    p = d.add_paragraph(); r = p.add_run()
    fld_begin = OxmlElement('w:fldChar'); fld_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve')
    instr.text = f'TOC \\o "1-{max_level}" \\h \\z \\u \\b {start_bmk} \\e {end_bmk}'
    fld_separate = OxmlElement('w:fldChar'); fld_separate.set(qn('w:fldCharType'), 'separate')
    fld_end = OxmlElement('w:fldChar'); fld_end.set(qn('w:fldCharType'), 'end')
    r._r.append(fld_begin); r._r.append(instr); r._r.append(fld_separate)
    p.add_run("<< Tabla de contenido (actualiza campos en Word) >>")
    r._r.append(fld_end)
    return d


def create_page_number_footer_doc(start_at: int = 1) -> Document:
    """Define inicio de numeración y pie con PAGE (el resto hereda)."""
    d = Document()
    sect = d.sections[0]
    sectPr = sect._sectPr
    pgNumType = sectPr.find(qn('w:pgNumType')) or OxmlElement('w:pgNumType')
    if pgNumType.getparent() is None:
        sectPr.append(pgNumType)
    pgNumType.set(qn('w:start'), str(start_at))
    # pie con PAGE
    footer = sect.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    r = p.add_run()
    fld_begin = OxmlElement('w:fldChar'); fld_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve'); instr.text = 'PAGE'
    fld_separate = OxmlElement('w:fldChar'); fld_separate.set(qn('w:fldCharType'), 'separate')
    fld_end = OxmlElement('w:fldChar'); fld_end.set(qn('w:fldCharType'), 'end')
    r._r.append(fld_begin); r._r.append(instr); r._r.append(fld_separate)
    p.add_run("1")
    r._r.append(fld_end)
    return d


# ---------------------- COMPOSICIÓN (AGRUPADA, SIN TÍTULOS ARTIFICIALES) ---------------------- #

def compose_full_document(
    input_path: str,
    output_path: str,
    sections: List[Dict],
    selected_indices_in_order: List[int],
    insert_toc: bool,
    toc_levels: List[int]
) -> None:
    """
    Genera: [Portada] -> [TOC opcional] -> [Numeración] -> [Padres reales (1 vez)] + [Secciones recortadas]
    Todo lo copiado viene del documento original, conservando estilos.
    """

    # Sanitizar selección
    valid_sel = []
    for idx in selected_indices_in_order or []:
        if isinstance(idx, int) and 0 <= idx < len(sections):
            sec = sections[idx]
            if sec['end_index'] > sec['start_index']:
                valid_sel.append(idx)
    selected_indices_in_order = valid_sel

    doc_all = Document(input_path)
    n_par = len(doc_all.paragraphs)

    # Portada = antes del primer heading target (si hay)
    first_section_start = min((s["start_index"] for s in sections), default=0)
    has_portada = first_section_start > 0

    base = Document()
    composer = Composer(base)
    temp_dirs = []

    # ---- 1) PORTADA ----
    if has_portada:
        slice_portada = slice_docx_by_paragraph_range(input_path, 0, first_section_start)
        temp_dirs.append(os.path.dirname(slice_portada))
        composer.append(Document(slice_portada))
    else:
        cover = Document(); cover.add_paragraph(""); composer.append(cover)

    # Salto de página tras portada
    sep = Document(); sep.add_paragraph().add_run().add_break(WD_BREAK.PAGE); composer.append(sep)

    # ---- 2) TOC opcional (excluye portada y el propio TOC) ----
    start_bmk = "TOCSTART"
    end_bmk = "TOCEND"
    if insert_toc:
        toc_doc = create_toc_doc_ranged(
            start_bmk, end_bmk, max_level=(max(toc_levels) if toc_levels else 3), title="Índice"
        )
        composer.append(toc_doc)
        # Salto de página tras TOC
        sep2 = Document(); sep2.add_paragraph().add_run().add_break(WD_BREAK.PAGE); composer.append(sep2)
        # Marcador de INICIO del rango del TOC (después del TOC)
        composer.append(create_bookmark_doc(start_bmk, start=True))

    # ---- 3) Arranque de numeración en 1 ----
    composer.append(create_page_number_footer_doc(start_at=1))

    # ---- 4) CUERPO AGRUPADO (solo contenido real del DOC) ----
    if selected_indices_in_order:
        emitted_parent_at_index = set()
        for i, idx in enumerate(selected_indices_in_order):
            sec = sections[idx]

            # Inserta cadena de padres: pega el párrafo del heading real UNA sola vez
            for parent in sec.get("parents", []):
                key = (parent["level"], parent["index"])
                if key not in emitted_parent_at_index:
                    p_slice = slice_docx_by_paragraph_range(input_path, parent["index"], parent["index"] + 1)
                    temp_dirs.append(os.path.dirname(p_slice))
                    composer.append(Document(p_slice))
                    emitted_parent_at_index.add(key)

            # Inserta la sección completa (incluye su propio heading real + contenido)
            s_slice = slice_docx_by_paragraph_range(input_path, sec['start_index'], sec['end_index'])
            temp_dirs.append(os.path.dirname(s_slice))
            composer.append(Document(s_slice))

            # Salto de página entre secciones seleccionadas
            if i < len(selected_indices_in_order) - 1:
                sep3 = Document(); sep3.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
                composer.append(sep3)
    else:
        # Si no se seleccionó nada, mete todo desde el primer target al final
        core_start = first_section_start if has_portada else 0
        if n_par > core_start:
            s_slice = slice_docx_by_paragraph_range(input_path, core_start, n_par)
            temp_dirs.append(os.path.dirname(s_slice))
            composer.append(Document(s_slice))

    # ---- 5) Marcador de FIN del rango del TOC ----
    if insert_toc:
        composer.append(create_bookmark_doc(end_bmk, start=False))

    # Guardar y limpiar
    composer.save(output_path)
    for d in temp_dirs:
        try:
            shutil.rmtree(d, ignore_errors=True)
        except Exception:
            pass


# ------------------------------ SANITIZACIÓN SELECCIÓN ------------------------------ #

def sanitize_selection(selected_indices, sections):
    """Devuelve solo índices válidos (y únicos) para la lista actual de secciones."""
    valid, seen, n = [], set(), len(sections)
    for idx in selected_indices or []:
        if isinstance(idx, int) and 0 <= idx < n and sections[idx]['end_index'] > sections[idx]['start_index']:
            if idx not in seen:
                valid.append(idx); seen.add(idx)
    return valid


# ------------------------------ INTERFAZ GUI ------------------------------ #

class ProgressDialog(tk.Toplevel):
    def __init__(self, parent, title="Progreso", mode="indeterminate"):
        super().__init__(parent)
        self.title(title)
        self.resizable(False, False)
        self.progress = ttk.Progressbar(self, mode=mode, length=360)
        self.label = ttk.Label(self, text="...")
        self.progress.grid(row=0, column=0, padx=16, pady=(16, 6))
        self.label.grid(row=1, column=0, padx=16, pady=(0, 12))
        self.grab_set()
        self.protocol("WM_DELETE_WINDOW", lambda: None)
        self.update_idletasks()
        w, h = self.winfo_width(), self.winfo_height()
        x = self.winfo_screenwidth() // 2 - w // 2
        y = self.winfo_screenheight() // 2 - h // 2
        self.geometry(f"+{x}+{y}")

    def set_text(self, text: str):
        self.label.config(text=text); self.update_idletasks()

    def set_fraction(self, frac: float):
        self.progress.config(mode="determinate", maximum=100, value=int(frac * 100))
        self.update_idletasks()

    def start(self, interval=10):
        self.progress.start(interval)

    def stop(self):
        self.progress.stop()


class App(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("Extractor DOCX — portada, numeración tras portada, agrupación + TOC opcional")
        self.geometry("1120x680"); self.minsize(980, 560)

        self.input_path: Optional[str] = None
        self.sections: List[Dict] = []
        self.filtered_indices: List[int] = []
        self.selected_indices: List[int] = []

        self.insert_toc_var = tk.BooleanVar(value=False)

        self.build_ui()

    def build_ui(self):
        top = ttk.Frame(self, padding=10); top.pack(side=tk.TOP, fill=tk.X)
        ttk.Button(top, text="📂 Cargar DOCX…", command=self.load_docx).pack(side=tk.LEFT)

        ttk.Label(top, text="  Niveles:").pack(side=tk.LEFT, padx=(12, 4))
        self.levels_var = tk.StringVar(value="1,2,3,4,5")
        ttk.Entry(top, width=16, textvariable=self.levels_var).pack(side=tk.LEFT)

        ttk.Checkbutton(top, text="Insertar Índice (TOC) tras la portada",
                        variable=self.insert_toc_var).pack(side=tk.LEFT, padx=(16, 4))

        ttk.Button(top, text="🔎 Detectar secciones", command=self.detect_and_list).pack(side=tk.LEFT, padx=8)
        ttk.Button(top, text="💾 Exportar…", command=self.export_selected).pack(side=tk.RIGHT)

        path_frame = ttk.Frame(self, padding=(10, 0)); path_frame.pack(side=tk.TOP, fill=tk.X)
        self.path_label = ttk.Label(path_frame, text="Archivo: (ninguno)", foreground="#555"); self.path_label.pack(side=tk.LEFT)

        mid = ttk.Frame(self, padding=10); mid.pack(side=tk.TOP, fill=tk.BOTH, expand=True)

        left = ttk.Frame(mid); left.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        ttk.Label(left, text="Disponibles").pack(anchor="w")
        filter_bar = ttk.Frame(left); filter_bar.pack(fill=tk.X, pady=(4, 6))
        ttk.Label(filter_bar, text="Buscar:").pack(side=tk.LEFT)
        self.search_var = tk.StringVar(); search_entry = ttk.Entry(filter_bar, textvariable=self.search_var)
        search_entry.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(6, 6))
        search_entry.bind("<KeyRelease>", lambda e: self.apply_filter())

        self.list_available = tk.Listbox(left, selectmode=tk.EXTENDED, activestyle='none', font=("Segoe UI", 10))
        self.list_available.pack(fill=tk.BOTH, expand=True)
        self.list_available.bind("<Double-Button-1>", lambda e: self.add_selected())

        center_btns = ttk.Frame(mid, padding=10); center_btns.pack(side=tk.LEFT, fill=tk.Y)
        ttk.Button(center_btns, text="➕ Añadir →", command=self.add_selected).pack(pady=4, fill=tk.X)
        ttk.Button(center_btns, text="← Quitar", command=self.remove_selected).pack(pady=4, fill=tk.X)
        ttk.Button(center_btns, text="✅ Seleccionar todo", command=self.select_all_left).pack(pady=(16, 4), fill=tk.X)
        ttk.Button(center_btns, text="🗑️ Limpiar selección", command=self.clear_right).pack(pady=4, fill=tk.X)

        right = ttk.Frame(mid); right.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        ttk.Label(right, text="Seleccionadas (arriba → primero)").pack(anchor="w")
        self.list_chosen = tk.Listbox(right, selectmode=tk.EXTENDED, activestyle='none', font=("Segoe UI", 10))
        self.list_chosen.pack(fill=tk.BOTH, expand=True)
        self.list_chosen.bind("<Double-Button-1>", lambda e: self.remove_selected())

        order_bar = ttk.Frame(right); order_bar.pack(fill=tk.X, pady=(6, 0))
        ttk.Button(order_bar, text="⬆️ Subir", command=self.move_up).pack(side=tk.LEFT, padx=(0, 6))
        ttk.Button(order_bar, text="⬇️ Bajar", command=self.move_down).pack(side=tk.LEFT, padx=(0, 6))

        tips = (
            "Consejos:\n"
            "• No se insertan títulos artificiales: todo conserva los estilos del documento original.\n"
            "• La numeración empieza en 1 tras la portada. Puedes insertar un TOC opcional que excluye la portada y él mismo.\n"
            "• Si eliges varias herramientas bajo el mismo bloque, el título padre se copia una única vez.\n"
            "• Doble-clic: izquierda añade / derecha quita. El orden de la derecha es el orden final.\n"
        )
        bottom = ttk.Frame(self, padding=(10, 4)); bottom.pack(side=tk.BOTTOM, fill=tk.X)
        ttk.Label(bottom, text=tips, justify=tk.LEFT, foreground="#444").pack(side=tk.LEFT, fill=tk.X, expand=True)

        self.status = ttk.Label(self, text="Listo", anchor="w", relief="sunken"); self.status.pack(side=tk.BOTTOM, fill=tk.X)

    # ---- Helpers ---- #
    def set_status(self, msg: str):
        self.status.config(text=msg); self.update_idletasks()

    def ask_levels(self) -> List[int]:
        raw = self.levels_var.get().strip()
        if not raw:
            return [1]
        try:
            return [int(x.strip()) for x in raw.split(",") if x.strip()]
        except ValueError:
            messagebox.showerror("Niveles inválidos", "Usa números separados por coma, p.ej. 1,2,3")
            return [1]

    # ---- Actions ---- #
    def load_docx(self):
        initial_dir = leer_ultima_ruta()
        path = filedialog.askopenfilename(
            title="Selecciona un archivo DOCX",
            filetypes=[("Documentos Word", "*.docx")],
            initialdir=initial_dir if initial_dir else None,
        )
        if not path:
            return
        self.input_path = path
        self.path_label.config(text=f"Archivo: {os.path.basename(path)}")

        # Guardar la carpeta para la próxima vez
        guardar_ruta(os.path.dirname(path))

        self.sections = []
        self.filtered_indices = []
        self.selected_indices = []
        self.list_available.delete(0, tk.END)
        self.list_chosen.delete(0, tk.END)
        self.search_var.set("")
        self.set_status("DOCX cargado. Pulsa 'Detectar secciones'.")

    def detect_and_list(self):
        if not self.input_path:
            messagebox.showinfo("Falta archivo", "Primero carga un archivo DOCX."); return
        levels = self.ask_levels()

        prog = ProgressDialog(self, title="Escaneando documento…", mode="indeterminate")
        prog.set_text("Analizando encabezados y jerarquías…"); prog.start(12); self.update_idletasks()
        try:
            self.sections = detect_sections(self.input_path, levels)
        except Exception as e:
            prog.destroy(); messagebox.showerror("Error detectando secciones", str(e)); return
        finally:
            prog.stop(); prog.destroy()

        self.selected_indices = []  # limpia selección al re-detectar
        self.selected_indices = sanitize_selection(self.selected_indices, self.sections)
        self.apply_filter()
        self.set_status(f"Secciones detectadas: {len(self.sections)}")

    def apply_filter(self):
        self.list_available.delete(0, tk.END)
        term = self.search_var.get().strip().lower()
        self.filtered_indices = []
        for i, s in enumerate(self.sections):
            parents_str = " / ".join(p["title"] for p in s.get("parents", []) if p["title"])
            row = f"[{i:02d}] {parents_str + ' / ' if parents_str else ''}{s['title']}  (L{s['level']}, párrafos {s['start_index']}..{s['end_index'] - 1})"
            if not term or term in row.lower():
                self.list_available.insert(tk.END, row); self.filtered_indices.append(i)
        self.set_status(f"Disponibles: {len(self.filtered_indices)} (filtro='{term}')")

    def select_all_left(self):
        if self.list_available.size() == 0: return
        self.list_available.select_set(0, tk.END)
        self.set_status("Todas las visibles seleccionadas.")

    def add_selected(self):
        sel_vis = list(self.list_available.curselection())
        sel_real = [self.filtered_indices[i] for i in sel_vis] if sel_vis else self.filtered_indices[:]
        for idx in sel_real:
            if 0 <= idx < len(self.sections) and idx not in self.selected_indices:
                self.selected_indices.append(idx)
        self.selected_indices = sanitize_selection(self.selected_indices, self.sections)
        self.refresh_right(); self.list_available.selection_clear(0, tk.END)

    def remove_selected(self):
        sel = list(self.list_chosen.curselection())
        for pos in reversed(sel):
            if 0 <= pos < len(self.selected_indices):
                del self.selected_indices[pos]
        self.selected_indices = sanitize_selection(self.selected_indices, self.sections)
        self.refresh_right()

    def clear_right(self):
        self.selected_indices = []; self.refresh_right()

    def move_up(self):
        sel = list(self.list_chosen.curselection())
        if not sel: return
        for pos in sel:
            if pos == 0: continue
            self.selected_indices[pos-1], self.selected_indices[pos] = self.selected_indices[pos], self.selected_indices[pos-1]
        self.refresh_right(new_selection=[max(0, p-1) for p in sel])

    def move_down(self):
        sel = list(self.list_chosen.curselection())
        if not sel: return
        for pos in reversed(sel):
            if pos >= len(self.selected_indices) - 1: continue
            self.selected_indices[pos+1], self.selected_indices[pos] = self.selected_indices[pos], self.selected_indices[pos+1]
        self.refresh_right(new_selection=[min(len(self.selected_indices)-1, p+1) for p in sel])

    def refresh_right(self, new_selection: Optional[List[int]] = None):
        self.selected_indices = sanitize_selection(self.selected_indices, self.sections)
        self.list_chosen.delete(0, tk.END)
        for idx in self.selected_indices:
            s = self.sections[idx]
            parents_str = " / ".join(p["title"] for p in s.get("parents", []) if p["title"])
            row = f"[{idx:02d}] {parents_str + ' / ' if parents_str else ''}{s['title']} (L{s['level']})"
            self.list_chosen.insert(tk.END, row)
        if new_selection:
            for pos in new_selection:
                if 0 <= pos < self.list_chosen.size(): self.list_chosen.selection_set(pos)
        self.set_status(f"Seleccionadas: {len(self.selected_indices)}")

    def export_selected(self):
        if not self.input_path:
            messagebox.showinfo("Falta archivo", "Primero carga un archivo DOCX."); return
        if self.sections is None:
            messagebox.showinfo("Faltan secciones", "Pulsa 'Detectar secciones' para listar."); return

        out_path = filedialog.asksaveasfilename(title="Guardar DOCX resultante", defaultextension=".docx",
                                                filetypes=[("Documentos Word", "*.docx")], initialfile="salida.docx")
        if not out_path: return

        if os.path.exists(out_path):
            try:
                with open(out_path, "ab"): pass
            except PermissionError:
                messagebox.showerror("Archivo en uso", "Cierra el archivo de salida en Word e inténtalo de nuevo.")
                return

        # Sanea selección
        self.selected_indices = sanitize_selection(self.selected_indices, self.sections)
        toc_levels = self.ask_levels()
        insert_toc = bool(self.insert_toc_var.get())

        # Pasos de progreso (añadimos 1 si hay TOC)
        N = max(1, len(self.selected_indices))
        total_steps = 2 + (1 if insert_toc else 0) + N  # portada + [TOC] + numeración + secciones
        step = 0
        prog = ProgressDialog(self, title="Exportando…", mode="determinate")
        prog.set_text("Preparando…"); prog.set_fraction(0.0)

        def tick(text):
            nonlocal step
            step += 1; prog.set_text(text); prog.set_fraction(min(1.0, step / total_steps))

        self.config(cursor="wait"); self.update_idletasks()
        try:
            tick("Añadiendo portada…")
            if insert_toc:
                tick("Insertando Índice (TOC)…")
            tick("Preparando numeración del cuerpo…")
            if self.selected_indices:
                for i, _ in enumerate(self.selected_indices, 1):
                    tick(f"Añadiendo sección {i}/{len(self.selected_indices)}…")
            else:
                tick("Añadiendo contenido…")

            compose_full_document(
                input_path=self.input_path,
                output_path=out_path,
                sections=self.sections,
                selected_indices_in_order=self.selected_indices,
                insert_toc=insert_toc,
                toc_levels=toc_levels
            )
        except Exception as e:
            self.config(cursor=""); prog.destroy()
            tb = traceback.format_exc()
            messagebox.showerror("Error exportando", f"{e}\n\nDetalle técnico:\n{tb}")
            return
        finally:
            self.config(cursor="")

        prog.set_text("Finalizando…"); prog.set_fraction(1.0); prog.after(300, prog.destroy)
        self.set_status(f"¡Listo! Guardado en: {out_path}")
        if insert_toc:
            messagebox.showinfo(
                "Completado",
                "Documento generado.\n\nAl abrir en Word, acepta **Actualizar los campos** para refrescar el Índice."
            )
        else:
            messagebox.showinfo("Completado", "Documento generado.")

def main():
    app = App(); app.mainloop()

if __name__ == "__main__":
    main()
